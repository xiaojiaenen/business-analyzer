#!/usr/bin/env python3
"""Source → Markdown extraction helper for the Beautiful Article skill.

Mechanically extracts text from a PDF / DOCX / HTML file (or URL) into rough
Markdown. It does NOT clean editorial noise, mark image placeholders, or record
extraction risk — that judgement stays with the agent (see Phase 1 +
references/source-to-markdown.md). The agent should review and refine the output
into source/source.md and write source/extraction-notes.md.

Usage:
    python3 source-to-markdown.py <input.pdf|.docx|.html|.htm|.txt|.md|URL> [-o out.md]

Dependencies are probed at runtime and optional. If a parser is missing the
script prints an install hint and degrades gracefully.
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path
from urllib.parse import urlparse


def _hint(pkg: str) -> str:
    return f"  (缺少 {pkg}，可安装：pip install {pkg})"


def from_pdf(path: str) -> str:
    try:
        import pdfplumber  # type: ignore
    except Exception:
        try:
            from pdfminer.high_level import extract_text  # type: ignore

            return extract_text(path)
        except Exception:
            print(_hint("pdfplumber 或 pdfminer.six"), file=sys.stderr)
            raise SystemExit(2)
    out = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            out.append(f"\n<!-- page {i} -->\n")
            out.append(page.extract_text() or "")
            for t in page.extract_tables() or []:
                out.append("\n" + _table_to_md(t) + "\n")
    return "\n".join(out)


def from_docx(path: str) -> str:
    try:
        import docx  # type: ignore
    except Exception:
        print(_hint("python-docx"), file=sys.stderr)
        raise SystemExit(2)
    doc = docx.Document(path)
    out = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            out.append("")
            continue
        style = (p.style.name or "").lower()
        if style.startswith("heading"):
            level = "".join(c for c in style if c.isdigit()) or "1"
            out.append("#" * min(int(level), 6) + " " + text)
        else:
            out.append(text)
    for table in doc.tables:
        rows = [[c.text.strip() for c in r.cells] for r in table.rows]
        if rows:
            out.append("\n" + _table_to_md(rows) + "\n")
    return "\n".join(out)


def from_html(html: str) -> str:
    try:
        from bs4 import BeautifulSoup  # type: ignore
    except Exception:
        print(_hint("beautifulsoup4"), file=sys.stderr)
        # crude fallback: strip tags
        import re

        return re.sub(r"<[^>]+>", "", html)
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "aside", "header"]):
        tag.decompose()
    main = soup.find("article") or soup.find("main") or soup.body or soup
    lines = []
    for el in main.find_all(
        ["h1", "h2", "h3", "h4", "p", "li", "pre", "blockquote"]
    ):
        text = el.get_text(" ", strip=True)
        if not text:
            continue
        name = el.name
        if name.startswith("h") and name[1:].isdigit():
            lines.append("#" * int(name[1:]) + " " + text)
        elif name == "li":
            lines.append("- " + text)
        elif name == "pre":
            lines.append("```\n" + text + "\n```")
        elif name == "blockquote":
            lines.append("> " + text)
        else:
            lines.append(text)
    return "\n\n".join(lines)


def from_url(url: str) -> str:
    try:
        import urllib.request

        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=30) as resp:  # noqa: S310
            html = resp.read().decode("utf-8", "replace")
    except Exception as e:  # pragma: no cover
        print(f"✗ 抓取失败：{e}", file=sys.stderr)
        print("  也可以用 agent 的网页抓取能力获取正文后再清理。", file=sys.stderr)
        raise SystemExit(2)
    return from_html(html)


def _table_to_md(rows) -> str:
    rows = [[("" if c is None else str(c)).replace("\n", " ").strip() for c in r] for r in rows]
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    head = "| " + " | ".join(rows[0]) + " |"
    sep = "| " + " | ".join(["---"] * width) + " |"
    body = ["| " + " | ".join(r) + " |" for r in rows[1:]]
    return "\n".join([head, sep, *body])


def main() -> None:
    ap = argparse.ArgumentParser(description="Source → Markdown extraction helper")
    ap.add_argument("input", help="PDF / DOCX / HTML / TXT / MD file or URL")
    ap.add_argument("-o", "--output", help="write Markdown here (default: stdout)")
    args = ap.parse_args()

    src = args.input
    parsed = urlparse(src)
    if parsed.scheme in ("http", "https"):
        md = from_url(src)
    else:
        p = Path(src)
        if not p.exists():
            print(f"✗ 文件不存在：{src}", file=sys.stderr)
            raise SystemExit(1)
        ext = p.suffix.lower()
        if ext == ".pdf":
            md = from_pdf(str(p))
        elif ext == ".docx":
            md = from_docx(str(p))
        elif ext in (".html", ".htm"):
            md = from_html(p.read_text("utf-8", "replace"))
        elif ext in (".md", ".markdown", ".txt"):
            md = p.read_text("utf-8", "replace")
        else:
            print(f"✗ 不支持的输入类型：{ext}", file=sys.stderr)
            raise SystemExit(1)

    md = (md or "").strip() + "\n"
    if args.output:
        out = Path(args.output)
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text(md, "utf-8")
        print(f"✓ 写入 {out}（{len(md)} 字符）。请人工清理噪音并补 extraction-notes.md。", file=sys.stderr)
    else:
        sys.stdout.write(md)


if __name__ == "__main__":
    main()
